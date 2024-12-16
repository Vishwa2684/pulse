from django.shortcuts import redirect
from django.http import HttpResponse, JsonResponse
from django.core.mail import EmailMessage
from django.conf import settings
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .serializers import GeminiSerializer
import google.generativeai as genai
from google.cloud import vision
import io
import speech_recognition as sr
from io import BytesIO
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv

load_dotenv()

# Google API key for Generative AI
API_KEY = "AIzaSyDGwSA4vpgACT1DzT7LBsuXryx5U3zNmGY"  
genai.configure(api_key=API_KEY)

# Basic Route
def index(request):
    return HttpResponse({'message': 'Hello, world! This is your Django app.'})

# Function to handle image recognition using Google's Vision API
def analyze_image(image_file):
    client = vision.ImageAnnotatorClient()

    with io.open(image_file, 'rb') as image:
        content = image.read()

    image = vision.Image(content=content)
    response = client.label_detection(image=image)
    labels = response.label_annotations

    return [label.description for label in labels]

# Function to handle speech recognition using SpeechRecognition library
def transcribe_audio(audio_file):
    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_file) as source:
        audio = recognizer.record(source)

    try:
        text = recognizer.recognize_google(audio)
        return text
    except sr.UnknownValueError:
        return "Could not understand audio"
    except sr.RequestError as e:
        return f"Error with the speech recognition service: {e}"

# Function to create a document (Word or PDF) from the given topic
def create_document(topic, doc_type):
    topic = "".join(c for c in topic if c.isalnum() or c.isspace()).strip()
    filename = f"{topic.replace(' ', '_')}.{doc_type}"
    content = BytesIO()

    if doc_type == "word":
        doc = Document()
        doc.add_heading(f"Report on {topic}", level=1)
        doc.add_heading("Introduction", level=2)
        doc.add_paragraph(f"This document provides an overview of the topic: {topic}.")
        doc.add_heading("Main Points", level=2)
        doc.add_heading("1. Overview", level=3)
        doc.add_paragraph(f"The topic {topic} encompasses various areas including ...")
        doc.add_heading("2. Key Challenges", level=3)
        doc.add_paragraph("While exploring this topic, some challenges include ...")
        doc.add_heading("Conclusion", level=2)
        doc.add_paragraph(f"In conclusion, {topic} is an important area for further study.")
        doc.save(content)
        content.seek(0)

    elif doc_type == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=f"Report on {topic}", ln=True, align="C")
        pdf.ln(10)
        pdf.multi_cell(0, 10, f"This document provides an overview of the topic: {topic}.")
        content.write(pdf.output(dest='S').encode('latin1'))
        content.seek(0)

    return content, filename


# API view to generate document and send via email
class DocumentGenerationView(APIView):
    def post(self, request):
        topic = request.data.get('topic')
        doc_type = request.data.get('doc_type', 'pdf').lower()
        recipient_email = request.data.get('email')

        if not topic or not recipient_email:
            return Response({"error": "Topic and recipient email are required."}, status=status.HTTP_400_BAD_REQUEST)

        if doc_type not in ["word", "pdf"]:
            return Response({"error": "Invalid document type. Supported types are 'word' and 'pdf'."},
                            status=status.HTTP_400_BAD_REQUEST)

        try:
            content, filename = create_document(topic, doc_type)

            email_subject = f"Generated Document: {filename}"
            email_body = f"Please find attached the document on {topic}."
            email = EmailMessage(subject=email_subject, body=email_body, to=[recipient_email])
            email.attach(filename, content.getvalue(), 'application/octet-stream')

            email.send()

            return Response({"message": f"Document '{filename}' sent successfully to {recipient_email}."},
                            status=status.HTTP_200_OK)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

# Gemini AI view to handle both text, image and audio input
class GeminiViewSet(APIView):
    def post(self, request):
        serializer = GeminiSerializer(data=request.data)
        if serializer.is_valid():
            prompt = serializer.validated_data['message']
            chat_history = serializer.validated_data.get('chatHistory', [])

            # Handle image file
            image_file = request.FILES.get('image')
            image_labels = analyze_image(image_file) if image_file else []

            # Handle audio file
            audio_file = request.FILES.get('audio')
            audio_text = transcribe_audio(audio_file) if audio_file else ""

            try:
                model = genai.GenerativeModel("gemini-pro")
                chat = model.start_chat(history=chat_history)

                # Incorporate image labels and audio text into the prompt
                if image_labels:
                    prompt += f"\nImage contains: {', '.join(image_labels)}."
                if audio_text:
                    prompt += f"\nVoice command: {audio_text}."

                response = chat.send_message(prompt)

                return Response({'response': response.text}, status=status.HTTP_200_OK)
            except Exception as e:
                return Response({'error': f'Error: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# Test view for simple "hello world"
class TestView(APIView):
    def post(self, request):
        return Response({"message": "hello world"})

# Define additional routes like login, logout if needed
def main(request):
    return HttpResponse({'message': 'Main route'})

def login(request):
    return HttpResponse({'message': 'Login route'})

def logout(request):
    return HttpResponse({'message': 'Logout route'})
